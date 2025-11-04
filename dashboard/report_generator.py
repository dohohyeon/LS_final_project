# ---------------------------------------------------
# 전력 데이터 통합 분석 보고서 생성 (폰트 통일 + 자동 삭제 + 최신 기능 반영)
# ---------------------------------------------------

import os, io, threading, time
import pandas as pd
import numpy as np
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# =========================
# ✅ 폰트 및 그래프 설정
# =========================
os.environ.setdefault("FONTCONFIG_PATH", "/usr/share/fonts/truetype/nanum")
WORD_FONT = "맑은 고딕"
FONT_NAME = "Noto Sans KR, Malgun Gothic, Apple SD Gothic Neo, sans-serif"

from matplotlib import rcParams
rcParams["font.family"] = ["Noto Sans KR", "Malgun Gothic", "Apple SD Gothic Neo", "sans-serif"]
rcParams["axes.unicode_minus"] = False

import plotly.io as pio
pio.templates.default = "plotly_white"
pio.defaults.font = dict(family=FONT_NAME, size=12, color="#222")

# =========================
# Plotly 폰트 강제 적용 (Kaleido 저장 직전)
# =========================
def _apply_korean_font(fig, family=FONT_NAME):
    fig.layout.font.family = family
    fig.update_layout(font=dict(family=family))
    for ax in [a for a in fig.layout if str(a).startswith("xaxis") or str(a).startswith("yaxis")]:
        axis = fig.layout[ax]
        if hasattr(axis, "title") and axis.title:
            axis.title.font = dict(family=family, size=12)
        if hasattr(axis, "tickfont") and axis.tickfont:
            axis.tickfont.family = family
    if hasattr(fig.layout, "legend") and fig.layout.legend:
        fig.layout.legend.font = dict(family=family, size=11)
    if getattr(fig.layout, "annotations", None):
        for a in fig.layout.annotations:
            a.font = dict(family=family, size=12)

# =========================
# 공통 표 스타일
# =========================
def format_table_uniform(table):
    table.style = "Table Grid"
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
            )
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = WORD_FONT
                    try:
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
                    except Exception:
                        pass
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    if ridx == 0:
                        r.font.bold = True


# =========================
# 메인 보고서 함수
# =========================
def generate_analysis_report(df, filtered_df, output_path="./reports/tab2_report.docx"):
    df, filtered_df = df.copy(), filtered_df.copy()
    df["측정일시"] = pd.to_datetime(df["측정일시"], errors="coerce")
    filtered_df["측정일시"] = pd.to_datetime(filtered_df["측정일시"], errors="coerce")
    filtered_df["시간"] = filtered_df["측정일시"].dt.hour

    # 수치형 변환
    for col in ["전력사용량(kWh)", "전기요금(원)", "수요전력(kW)", "지상역률(%)", "진상역률(%)"]:
        for d in [df, filtered_df]:
            if col in d.columns:
                d[col] = pd.to_numeric(d[col], errors="coerce")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = WORD_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)

    # ✅ 페이지 테두리 추가 (복원)
    sectPr = doc.sections[0]._sectPr
    sectPr.append(parse_xml('''
        <w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:offsetFrom="page">
            <w:top w:val="single" w:sz="12" w:space="24" w:color="auto"/>
            <w:left w:val="single" w:sz="12" w:space="24" w:color="auto"/>
            <w:bottom w:val="single" w:sz="12" w:space="24" w:color="auto"/>
            <w:right w:val="single" w:sz="12" w:space="24" w:color="auto"/>
        </w:pgBorders>
    '''))

    # 제목
    p = doc.add_paragraph("전력 데이터 통합 분석 보고서")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(18)
    p.runs[0].bold = True

    sd, ed = filtered_df["측정일시"].min(), filtered_df["측정일시"].max()
    doc.add_paragraph(f"분석 기간: {sd:%Y-%m-%d} ~ {ed:%Y-%m-%d}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"생성일: {datetime.now():%Y-%m-%d %H:%M}  |  작성자: 관리자").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1️⃣ 주요 지표 요약
    doc.add_heading("1. 주요 지표 요약", level=1)

    # ✅ 기본 통계 (현재 구간)
    total_usage = filtered_df["전력사용량(kWh)"].sum()
    total_cost  = filtered_df["전기요금(원)"].sum()
    avg_price   = total_cost / total_usage if total_usage > 0 else 0
    carbon      = total_usage * 0.000331

    # ✅ 전월 동일기간 구하기 (대시보드와 동일)
    sd, ed = filtered_df["측정일시"].min(), filtered_df["측정일시"].max()
    prev_start = sd - pd.DateOffset(months=1)
    prev_end   = ed - pd.DateOffset(months=1)
    prev_df = df[(df["측정일시"] >= prev_start) & (df["측정일시"] <= prev_end)]

    # ✅ 이전 기간 통계
    if not prev_df.empty:
        prev_usage  = prev_df["전력사용량(kWh)"].sum()
        prev_cost   = prev_df["전기요금(원)"].sum()
        prev_price  = prev_cost / prev_usage if prev_usage > 0 else 0
        prev_carbon = prev_usage * 0.000331
    else:
        prev_usage = prev_cost = prev_price = prev_carbon = np.nan

    # ✅ 증감률 함수 (대시보드와 동일)
    def pct_change(curr, prev):
        if pd.isna(prev) or prev == 0:
            return None
        return (curr - prev) / prev * 100

    usage_diff  = pct_change(total_usage, prev_usage)
    cost_diff   = pct_change(total_cost, prev_cost)
    price_diff  = pct_change(avg_price, prev_price)
    carbon_diff = pct_change(carbon, prev_carbon)

    # ✅ 테이블 생성
    t = doc.add_table(rows=2, cols=4)
    headers = ["전력사용량(kWh)", "전기요금(원)", "평균단가(원/kWh)", "탄소배출량(tCO₂)"]
    vals = [f"{total_usage:,.1f}", f"{total_cost:,.0f}", f"{avg_price:,.2f}", f"{carbon:,.3f}"]

    def fmt_diff(v):
        if v is None:
            return "-"
        sign = "+" if v > 0 else ""
        return f"{sign}{v:.1f}%"

    changes = [fmt_diff(usage_diff), fmt_diff(cost_diff), fmt_diff(price_diff), fmt_diff(carbon_diff)]

    for i, h in enumerate(headers):
        t.cell(0, i).text = h
        t.cell(1, i).text = f"{vals[i]}\n({changes[i]})"
    format_table_uniform(t)

    # 설명문
    doc.add_paragraph(
        f"- {sd:%Y-%m-%d} ~ {ed:%Y-%m-%d} 구간의 주요 지표와 전월 동일기간({prev_start:%Y-%m-%d} ~ {prev_end:%Y-%m-%d}) 대비 증감률을 표시합니다. \n"
        "- 전월 대비 증감률을 함께 제시하여 에너지 사용 추세를 파악할 수 있습니다.\n "
        "- 평균 단가의 상승은 피크 부하 시간대 사용량 증가 또는 요금제 변경의 영향을 시사할 수 있습니다."
    )
    doc.add_paragraph()

    # 2️⃣ 요일·시간대별 전력 사용 패턴
    doc.add_heading("2. 요일·시간대별 전력 사용 패턴", level=1)
    weekday_map = {"Monday": "월요일", "Tuesday": "화요일", "Wednesday": "수요일",
                   "Thursday": "목요일", "Friday": "금요일", "Saturday": "토요일", "Sunday": "일요일"}
    filtered_df["요일"] = filtered_df["측정일시"].dt.day_name().map(weekday_map)

    wd = (filtered_df.groupby("요일")["전력사용량(kWh)"].mean()
          .reindex(["월요일", "화요일", "수요일", "목요일", "금요일", "토요일", "일요일"]).reset_index())
    hr = (filtered_df.groupby("시간")["전력사용량(kWh)"].mean()
          .reset_index().sort_values("시간"))

    fig_pattern = make_subplots(rows=2, cols=1,
                                subplot_titles=("요일별 평균 전력사용량", "시간대별 평균 전력사용량"),
                                vertical_spacing=0.18)
    fig_pattern.add_bar(x=wd["요일"], y=wd["전력사용량(kWh)"], marker_color="#3A86FF", row=1, col=1)
    fig_pattern.add_trace(
        go.Scatter(x=hr["시간"].values, y=hr["전력사용량(kWh)"].values, mode="lines+markers",
                   line=dict(width=2), marker=dict(size=6)),
        row=2, col=1
    )
    fig_pattern.update_layout(showlegend=False, hovermode="x unified", template="plotly_white",
                              height=700, plot_bgcolor="#ffffff", paper_bgcolor="#ffffff")
    _apply_korean_font(fig_pattern)
    buf = io.BytesIO()
    fig_pattern.write_image(buf, format="png")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(6))
    doc.add_paragraph(
        "- 요일별 패턴을 보면 평일과 주말 간 전력 사용량의 차이를 확인할 수 있습니다. \n"
        "- 시간대별 그래프에서는 주간과 야간의 사용 패턴이 명확하게 구분되며, 특정 시간대(예: 오전 9시~11시, 오후 2시~5시)에 전력 수요가 집중되는 경향이 나타납니다. \n"
        "- 이는 생산 공정 또는 설비 가동 스케줄과 밀접한 관련이 있습니다."
    )
    doc.add_paragraph()

    # 3-1️⃣ 피크 수요 및 역률 분석 - 그래프 (수요전력 + 관리기준선 + 피크 표시)
    if "수요전력(kW)" in filtered_df.columns:
        dfp = filtered_df.copy()

        # ✅ 안전한 숫자 변환
        for c in ["수요전력(kW)", "관리기준선(kW)"]:
            if c in dfp.columns:
                dfp[c] = pd.to_numeric(dfp[c], errors="coerce")

        # ✅ 시간 문자열 변환 (렌더링 안정성 확보)
        dfp["측정일시"] = pd.to_datetime(dfp["측정일시"], errors="coerce")
        dfp = dfp.dropna(subset=["측정일시", "수요전력(kW)"])
        dfp = dfp.sort_values("측정일시")
        dfp["시간라벨"] = dfp["측정일시"].dt.strftime("%m-%d %H:%M")

        # ✅ 그래프 생성
        fig_peak = go.Figure()

        # 1️⃣ 수요전력선
        fig_peak.add_trace(go.Scatter(
            x=dfp["시간라벨"],
            y=dfp["수요전력(kW)"],
            mode="lines",
            name="수요전력(kW)",
            line=dict(width=2.2, color="#007bff")
        ))

        # 2️⃣ 관리기준선
        if "관리기준선(kW)" in dfp.columns and not dfp["관리기준선(kW)"].isna().all():
            fig_peak.add_trace(go.Scatter(
                x=dfp["시간라벨"],
                y=dfp["관리기준선(kW)"],
                mode="lines",
                name="관리기준선(kW)",
                line=dict(color="#999", width=2, dash="dash")
            ))

        # 3️⃣ 기준 초과 마커 (채워진 실색 원)
        if "관리기준선(kW)" in dfp.columns:
            exceed_df = dfp[dfp["수요전력(kW)"] > dfp["관리기준선(kW)"]]
            if not exceed_df.empty:
                fig_peak.add_trace(go.Scatter(
                    x=exceed_df["시간라벨"],
                    y=exceed_df["수요전력(kW)"],
                    mode="markers",
                    name="기준 초과",
                    marker=dict(
                        symbol="circle",
                        size=6,
                        color="#ff4d4d",       # 실색 원
                        line=dict(width=0)
                    )
                ))

        # 4️⃣ 최대 피크 (빨간 별)
        if not dfp.empty:
            peak_idx = dfp["수요전력(kW)"].idxmax()
            peak_point = dfp.loc[peak_idx]
            fig_peak.add_trace(go.Scatter(
                x=[peak_point["시간라벨"]],
                y=[peak_point["수요전력(kW)"]],
                mode="markers+text",
                text=["최대 피크"],
                textposition="top center",
                name="최대 피크",
                marker=dict(
                    symbol="star",
                    size=16,
                    color="#e60000",
                    line=dict(width=1, color="#660000")
                )
            ))

        # ✅ 축 및 레이아웃 (시간축 단순 문자열)
        fig_peak.update_layout(
            title="기간 내 전력 사용량 추이 (관리기준선 및 최대 피크 표시)",
            template="plotly_white",
            hovermode="x unified",
            plot_bgcolor="#fff",
            paper_bgcolor="#fff",
            showlegend=True,
            height=420,
            xaxis=dict(
                tickangle=45,
                showgrid=True,
                gridcolor="#eee",
                tickmode="auto",
                nticks=8
            )
        )

        # ✅ 폰트 적용 및 그림 삽입
        _apply_korean_font(fig_peak)
        buf = io.BytesIO()
        fig_peak.write_image(buf, format="png")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(6))

    doc.add_paragraph(
        "- 피크 수요 시간과 관리기준선 대비 초과 구간을 시각화하였습니다. \n"
        "- 초과 비율이 높을수록 부하 분산 및 피크 억제 대책 검토가 필요합니다.\n"
        "- 최고 피크 전력을 별표 마크로 표시하였습니다."
    )
    doc.add_paragraph()

    # 3-2️⃣ 역률 변화 추이
    doc.add_heading("3-2. 역률 변화 추이", level=1)
    pf_df = filtered_df.sort_values("측정일시").dropna(subset=["측정일시"])
    xvals = pf_df["측정일시"].dt.to_pydatetime()

    fig_pf = go.Figure()
    if "지상역률(%)" in pf_df:
        fig_pf.add_trace(go.Scatter(x=xvals, y=pf_df["지상역률(%)"],
                                    mode="lines", name="지상역률(%)",
                                    line=dict(color="#1f77b4", width=2)))
    if "진상역률(%)" in pf_df:
        fig_pf.add_trace(go.Scatter(x=xvals, y=pf_df["진상역률(%)"],
                                    mode="lines", name="진상역률(%)",
                                    line=dict(color="#ff7f0e", width=2)))

    fig_pf.add_hrect(y0=95, y1=100, fillcolor="rgba(40,167,69,0.15)", line_width=0)
    fig_pf.add_hrect(y0=0, y1=90, fillcolor="rgba(220,53,69,0.12)", line_width=0)
    fig_pf.add_hline(y=95, line_dash="dash", line_color="#28a745")
    fig_pf.add_hline(y=90, line_dash="dash", line_color="#dc3545")

    fig_pf.update_layout(title="역률 변화 추이 (지상/진상)",
                        template="plotly_white", height=400,
                        yaxis_title="역률(%)", xaxis_title="시간")
    fig_pf.update_xaxes(type="date", tickformat="%m-%d %H:%M", tickangle=45)
    _apply_korean_font(fig_pf)
    buf = io.BytesIO()
    fig_pf.write_image(buf, format="png")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(6))
    doc.add_paragraph(
        "- 역률 변화 추이를 통해 감면(95% 이상) 및 할증(90% 미만) 구간을 한눈에 파악할 수 있습니다. \n"
        "- 지상역률은 주간(09~23시), 진상역률은 야간(23~09시)에 주로 나타납니다."
    )
    doc.add_paragraph()

    # 3-3️⃣ 역률 적용 시간대별 작업유형 구성비 (지상 / 진상)
    doc.add_heading("3-3. 역률 적용 시간대별 작업유형 구성비 (지상 / 진상)", level=1)
    color_map = {"Light_Load": "#2ecc71", "Medium_Load": "#f39c12", "Maximum_Load": "#e74c3c"}

    for pf_label, cond in [
        ("지상역률 적용 구간 (09시~23시)", (filtered_df["시간"] >= 9) & (filtered_df["시간"] < 23)),
        ("진상역률 적용 구간 (23시~09시)", (filtered_df["시간"] >= 23) | (filtered_df["시간"] < 9))
    ]:
        sub_df = filtered_df[cond]
        if "작업유형" not in sub_df or sub_df.empty:
            continue

        # 시간대별 비율
        ratio_df = sub_df.groupby(["시간", "작업유형"]).size().reset_index(name="건수")
        ratio_df["비율(%)"] = ratio_df.groupby("시간")["건수"].transform(lambda x: x / x.sum() * 100)
        total_ratio = sub_df["작업유형"].value_counts(normalize=True).mul(100).reset_index()
        total_ratio.columns = ["작업유형", "비율(%)"]

        # 누적 막대
        fig_bar = px.bar(
            ratio_df, x="시간", y="비율(%)", color="작업유형",
            color_discrete_map=color_map,
            title=f"{pf_label} 시간대별 작업유형 비율 (누적 막대)",
            labels={"비율(%)": "작업유형 비율(%)"}
        )
        fig_bar.update_layout(barmode="stack", template="plotly_white",
                              height=420, xaxis=dict(dtick=1), hovermode="x unified")
        _apply_korean_font(fig_bar)
        buf = io.BytesIO()
        fig_bar.write_image(buf, format="png")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(5.5))

        # 도넛 차트
        fig_pie = px.pie(
            total_ratio, values="비율(%)", names="작업유형",
            title=f"{pf_label} 전체 구성비",
            color="작업유형",
            color_discrete_map=color_map, hole=0.45
        )
        fig_pie.update_layout(template="plotly_white", height=380)
        _apply_korean_font(fig_pie)
        buf = io.BytesIO()
        fig_pie.write_image(buf, format="png")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(4.5))
        doc.add_paragraph(
            f"- {pf_label} 구간에서 시간대별 작업유형 비율과 전체 구성비를 나타냅니다. \n"
            "- 빨간색(Maximum_Load)은 피크 부하, 주황색(Medium_Load)은 중간 부하, 초록색(Light_Load)은 저부하를 의미합니다."
        )
    doc.add_paragraph()

    # 4️⃣ 시간대별 작업유형별 전기요금 현황
    doc.add_heading("4. 시간대별 작업유형별 전기요금 현황", level=1)
    if "작업유형" in filtered_df.columns:
        g = (filtered_df.groupby(["시간", "작업유형"])["전기요금(원)"].sum().reset_index())
        color_map = {"Light_Load": "#2ecc71", "Medium_Load": "#f39c12", "Maximum_Load": "#e74c3c"}
        fig_cost = px.bar(
            g, x="시간", y="전기요금(원)", color="작업유형",
            title="시간대별 작업유형별 전기요금 현황 (누적 막대)",
            labels={"전기요금(원)": "전기요금(원)", "시간": "시간대"},
            color_discrete_map=color_map, text_auto=".2s"
        )
        fig_cost.update_layout(
            barmode="stack", template="plotly_white", hovermode="x unified",
            plot_bgcolor="#fff", paper_bgcolor="#fff",
            legend_title="작업유형", xaxis=dict(dtick=1), height=500
        )
        _apply_korean_font(fig_cost)
        buf = io.BytesIO()
        fig_cost.write_image(buf, format="png")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(6))
    doc.add_paragraph(
        "- 작업 유형별로 전기요금이 집중되는 시간대를 확인할 수 있습니다.\n "
        "- ‘Maximum_Load’ 구간이 특정 시간대에 편중되어 있다면 생산 일정 조정이나 부하 분산을 통해 요금 절감 효과를 기대할 수 있습니다."
    )
    doc.add_paragraph()

    # 5️⃣ 전력사용량 시계열 추이
    doc.add_heading("5. 전력사용량 시계열 추이", level=1)
    ts = (filtered_df[["측정일시", "전력사용량(kWh)"]]
          .dropna().drop_duplicates(subset=["측정일시"]).sort_values("측정일시"))
    xvals = ts["측정일시"].dt.to_pydatetime()
    fig_ts = go.Figure()
    fig_ts.add_trace(go.Scatter(x=xvals, y=ts["전력사용량(kWh)"],
                                mode="lines", line=dict(width=1.8, color="#007bff")))
    fig_ts.update_layout(title="전력사용량 시계열 추이", template="plotly_white",
                         hovermode="x unified", plot_bgcolor="#fff", paper_bgcolor="#fff")
    fig_ts.update_xaxes(type="date", tickformat="%m-%d", dtick=86400000*3,
                        tickangle=45, showgrid=True, gridcolor="#eee")
    _apply_korean_font(fig_ts)
    buf = io.BytesIO()
    fig_ts.write_image(buf, format="png")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(6))
    doc.add_paragraph(
        "- 시계열 추이를 통해 전력 사용의 계절적 변동이나 공정별 특이 패턴을 파악할 수 있습니다.\n "
        "- 특정 기간에 급격한 사용량 증가는 설비 점검, 신규 라인 가동, 또는 계절적 요인에 의한 것일 수 있습니다.\n "
        "- 이를 바탕으로 향후 수요 예측 및 전력 계약전력 조정에 활용할 수 있습니다."
    )
    doc.add_paragraph()

    # 8️⃣ 개선사항 및 제안 (통합형: 기존 조건 + 역률/작업유형 기반)
    doc.add_heading("8. 개선사항 및 제안", level=1)
    suggestions = []

    # ==========================================================
    # ✅ 주요 변수 재정의 (상단 주요지표 계산 방식과 독립적으로)
    # ==========================================================
    total_usage = filtered_df["전력사용량(kWh)"].sum() if "전력사용량(kWh)" in filtered_df else np.nan
    total_cost  = filtered_df["전기요금(원)"].sum() if "전기요금(원)" in filtered_df else np.nan
    avg_price   = total_cost / total_usage if total_usage and total_usage > 0 else np.nan

    # 전월 동일기간 계산 (대시보드 로직과 동일)
    if "측정일시" in filtered_df.columns:
        sd, ed = filtered_df["측정일시"].min(), filtered_df["측정일시"].max()
        prev_start = sd - pd.DateOffset(months=1)
        prev_end   = ed - pd.DateOffset(months=1)
        prev_df = df[(df["측정일시"] >= prev_start) & (df["측정일시"] <= prev_end)]
    else:
        prev_df = pd.DataFrame()

    if not prev_df.empty:
        pu = prev_df["전력사용량(kWh)"].sum()
        pc = prev_df["전기요금(원)"].sum()
        pp = pc / pu if pu > 0 else np.nan
    else:
        pu = pc = pp = np.nan

    # 피크 전력 및 평균 역률
    peak_kw = filtered_df["수요전력(kW)"].max() if "수요전력(kW)" in filtered_df else np.nan
    avg_lag = filtered_df["지상역률(%)"].mean() if "지상역률(%)" in filtered_df else np.nan
    avg_lead = filtered_df["진상역률(%)"].mean() if "진상역률(%)" in filtered_df else np.nan
    mean_pf = np.nanmean([avg_lag, avg_lead])

    # ==========================================================
    # ---------- 기존 조건 유지 ----------
    # ==========================================================
    # 전월 대비 전력사용량 증가
    if pu and not np.isnan(pu) and total_usage > pu * 1.1:
        suggestions.append(("에너지 절감", "전월 대비 사용량이 크게 증가하였습니다. 비효율 설비 점검 및 운전시간 조정이 필요합니다."))

    # 피크전력 높음
    if not np.isnan(peak_kw) and peak_kw > 500:
        suggestions.append(("피크 부하 관리", "피크 부하 시간대의 설비 분산 및 가동 최적화가 필요합니다."))

    # 평균 역률 낮음
    if not np.isnan(mean_pf) and mean_pf < 95:
        suggestions.append(("역률 개선", f"평균 역률이 {mean_pf:.1f}%로 낮게 나타났습니다. 역률 보상장치(콘덴서 등) 점검을 권장합니다."))

    # 평균 단가 상승
    if not np.isnan(pp) and not np.isnan(avg_price) and avg_price > pp * 1.05:
        suggestions.append(("요금 최적화", "시간대별 요금제 검토 및 계약전력 조정이 필요합니다."))

    # 작업유형별 불균형
    if "작업유형" in filtered_df.columns:
        g_ratio = filtered_df["작업유형"].value_counts(normalize=True)
        if "Maximum_Load" in g_ratio and g_ratio["Maximum_Load"] > 0.6:
            suggestions.append(("부하 분산", "‘Maximum_Load’ 구간이 전체의 60% 이상을 차지합니다. 생산 일정 재조정 및 부하 분산 검토가 필요합니다."))

    # ==========================================================
    # ---------- 새 조건 추가 (역률 기반) ----------
    # ==========================================================
    low_lag = (filtered_df["지상역률(%)"] < 90).mean() * 100 if "지상역률(%)" in filtered_df else 0
    low_lead = (filtered_df["진상역률(%)"] < 90).mean() * 100 if "진상역률(%)" in filtered_df else 0

    if low_lag > 20:
        suggestions.append(("지상역률 개선", f"지상역률이 90% 미만인 구간이 전체의 {low_lag:.1f}%입니다. 콘덴서 노후 또는 과부하 구간을 점검하십시오."))

    if low_lead > 20:
        suggestions.append(("진상역률 개선", f"진상역률이 90% 미만인 구간이 {low_lead:.1f}%입니다. 과보상 가능성이 있으므로 리액터 조정을 검토하십시오."))

    # ==========================================================
    # ---------- 작업유형별 저역률 분석 ----------
    # ==========================================================
    if "작업유형" in filtered_df.columns and "지상역률(%)" in filtered_df.columns:
        pf_type_df = filtered_df.assign(평균역률=filtered_df[["지상역률(%)", "진상역률(%)"]].mean(axis=1))
        pf_type_df["저역률"] = (pf_type_df["평균역률"] < 90).astype(int)
        low_pf_ratio = pf_type_df.groupby("작업유형")["저역률"].mean().mul(100).to_dict()
        if low_pf_ratio:
            worst_type = max(low_pf_ratio, key=low_pf_ratio.get)
            if low_pf_ratio[worst_type] > 30:
                suggestions.append(("작업유형별 개선",
                                    f"‘{worst_type}’ 구간의 저역률 비율이 {low_pf_ratio[worst_type]:.1f}%로 높습니다. "
                                    "피크 시간대 부하 분산 및 보상장치 점검을 권장합니다."))

    # ==========================================================
    # ---------- 단기 급등 진단 ----------
    # ==========================================================
    if "수요전력(kW)" in filtered_df.columns:
        df_peak = filtered_df.dropna(subset=["수요전력(kW)"]).sort_values("측정일시")
        if len(df_peak) > 20:
            last_peaks = df_peak["수요전력(kW)"].tail(3).mean()
            prev_peaks = df_peak["수요전력(kW)"].iloc[-10:-3].mean()
            if last_peaks > prev_peaks * 1.2:
                suggestions.append(("단기 부하 급증", "최근 피크전력이 단기간 내 20% 이상 상승하였습니다. 특정 설비 이상 또는 운전 패턴 변화를 점검하십시오."))

    # ==========================================================
    # ---------- 기본 제안 ----------
    # ==========================================================
    if not suggestions:
        suggestions = [
            ("운영 상태 양호", "전력 및 역률 상태가 안정적으로 유지되고 있습니다. 현재 운영 패턴을 지속하십시오."),
        ]

    # ==========================================================
    # ---------- 표 생성 ----------
    # ==========================================================
    t2 = doc.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text, t2.rows[0].cells[1].text = "개선 항목", "내용"
    for k, v in suggestions:
        r = t2.add_row().cells
        r[0].text, r[1].text = k, v
    format_table_uniform(t2)

    doc.add_paragraph(
        "\n- 위 개선 항목은 전력·요금·피크·역률 및 작업유형별 분석 결과를 종합적으로 고려하여 자동 산출되었습니다."
    )



    # === 폰트 일괄 적용 ===
    for style in doc.styles:
        if hasattr(style, "font"):
            style.font.name = WORD_FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)

    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = WORD_FONT
            try:
                r._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
            except Exception:
                pass

    # 저장 및 삭제 스케줄
    doc.save(output_path)
    def delayed_delete(path):
        time.sleep(10)
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass
    threading.Thread(target=delayed_delete, args=(output_path,), daemon=True).start()

    return output_path